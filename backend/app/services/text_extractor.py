"""Extract raw text from any resume format: PDF, DOCX, PNG, JPG, scanned PDFs.

This module handles text extraction only — the AI agent handles understanding.
Strategy:
1. PDF → pdfplumber (text-based PDFs)
2. PDF with no text → OCR fallback via PaddleOCR
3. DOCX → python-docx
4. Images (PNG/JPG) → PaddleOCR
"""

import io
import os
from pathlib import Path

import pdfplumber
from pypdf import PdfReader
from docx import Document


def _extract_text_from_pdf(file_path: Path) -> tuple[str, dict]:
    """Extract text from PDF. Falls back to OCR if text extraction yields little."""
    text_parts = []
    metadata = {}

    reader = PdfReader(str(file_path))
    metadata["page_count"] = len(reader.pages)

    # Try text-based extraction first
    try:
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception:
        # Fallback to pypdf
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    full_text = "\n".join(text_parts).strip()

    # If very little text extracted, try OCR (scanned PDF)
    if len(full_text) < 50:
        ocr_text = _ocr_pdf(file_path)
        if ocr_text and len(ocr_text) > len(full_text):
            full_text = ocr_text
            metadata["extraction_method"] = "ocr"
        else:
            metadata["extraction_method"] = "text_sparse"
    else:
        metadata["extraction_method"] = "text"

    return full_text, metadata


def _ocr_pdf(file_path: Path) -> str:
    """OCR a scanned PDF by converting pages to images."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(file_path))
        all_text = []

        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")
            text = _ocr_image_bytes(img_bytes)
            if text:
                all_text.append(text)

        doc.close()
        return "\n".join(all_text)
    except ImportError:
        return ""
    except Exception:
        return ""


def _ocr_image_bytes(img_bytes: bytes) -> str:
    """Run OCR on image bytes using PaddleOCR."""
    try:
        from paddleocr import PaddleOCR
        import numpy as np
        from PIL import Image

        ocr = PaddleOCR(use_angle_cls=True, lang="en", show_log=False)
        img = Image.open(io.BytesIO(img_bytes))
        img_array = np.array(img)

        result = ocr.ocr(img_array, cls=True)
        if not result or not result[0]:
            return ""

        lines = []
        for line in result[0]:
            text = line[1][0]
            if text.strip():
                lines.append(text.strip())

        return "\n".join(lines)
    except ImportError:
        return ""
    except Exception:
        return ""


def _extract_text_from_image(file_path: Path) -> tuple[str, dict]:
    """Extract text from image file using OCR."""
    metadata = {"extraction_method": "ocr", "page_count": 1}

    with open(file_path, "rb") as f:
        img_bytes = f.read()

    text = _ocr_image_bytes(img_bytes)
    return text, metadata


def _extract_text_from_docx(file_path: Path) -> tuple[str, dict]:
    """Extract text from DOCX file."""
    doc = Document(str(file_path))
    metadata = {"paragraph_count": len(doc.paragraphs), "extraction_method": "docx"}

    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))

    if doc.tables:
        metadata["tables_found"] = len(doc.tables)

    return "\n".join(text_parts), metadata


def extract_text(file_path: Path) -> tuple[str, dict]:
    """Extract text from any supported file format.

    Returns (raw_text, metadata_dict).
    Supported: .pdf, .docx, .doc, .png, .jpg, .jpeg
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _extract_text_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return _extract_text_from_docx(file_path)
    elif suffix in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"):
        return _extract_text_from_image(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
