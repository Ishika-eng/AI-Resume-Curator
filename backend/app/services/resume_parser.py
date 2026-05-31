import re
from pathlib import Path

import pdfplumber
from docx import Document
from pypdf import PdfReader

from app.models.resume import ParsedResume, ParsedSection

SECTION_PATTERNS = [
    r"(?i)^(summary|professional\s+summary|objective|profile|about\s+me)",
    r"(?i)^(experience|work\s+experience|employment|professional\s+experience)",
    r"(?i)^(education|academic|qualifications)",
    r"(?i)^(skills|technical\s+skills|core\s+competencies|technologies)",
    r"(?i)^(projects|personal\s+projects|academic\s+projects)",
    r"(?i)^(certifications?|licenses?|credentials)",
    r"(?i)^(achievements?|awards?|honors?)",
    r"(?i)^(publications?|research)",
    r"(?i)^(volunteer|community|extracurricular)",
    r"(?i)^(interests?|hobbies)",
    r"(?i)^(references?)",
    r"(?i)^(contact|personal\s+information)",
]


def _is_section_header(line: str) -> bool:
    cleaned = line.strip().rstrip(":")
    if not cleaned or len(cleaned) > 60:
        return False
    return any(re.match(p, cleaned) for p in SECTION_PATTERNS)


def _extract_section_title(line: str) -> str:
    return line.strip().rstrip(":").title()


def _clean_content_lines(lines: list[str]) -> list[str]:
    """Remove very short fragments and merge broken lines from column extraction."""
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Skip very short orphan fragments (likely column artifacts)
        # unless they look like skills or meaningful short items
        if len(stripped) < 3 and not stripped.isalpha():
            continue
        cleaned.append(stripped)
    return cleaned


def _segment_into_sections(text: str) -> list[ParsedSection]:
    lines = text.split("\n")
    sections: list[ParsedSection] = []
    current_title = "Header"
    current_lines: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Handle cases where two section headers appear on the same line
        # e.g., "PROJECTS SKILLS" from two-column layouts
        if _is_section_header(stripped):
            # Check if this line contains multiple headers
            words = stripped.split()
            multi_headers = []
            for w in words:
                if _is_section_header(w):
                    multi_headers.append(w)

            if len(multi_headers) >= 2:
                # Save current section
                if current_lines:
                    sections.append(
                        ParsedSection(title=current_title, content=_clean_content_lines(current_lines))
                    )
                # Start with the first header, second will come as separate section
                current_title = _extract_section_title(multi_headers[0])
                current_lines = []
                # We'll pick up the second header on its own in subsequent content
                continue

            if current_lines:
                sections.append(
                    ParsedSection(title=current_title, content=_clean_content_lines(current_lines))
                )
            current_title = _extract_section_title(stripped)
            current_lines = []
        else:
            current_lines.append(stripped)

    if current_lines:
        sections.append(ParsedSection(title=current_title, content=_clean_content_lines(current_lines)))

    return sections


def _extract_text_from_pdf(file_path: Path) -> tuple[str, dict]:
    text_parts = []
    metadata = {}

    reader = PdfReader(str(file_path))
    metadata["page_count"] = len(reader.pages)
    pdf_meta = reader.metadata
    if pdf_meta:
        metadata["pdf_metadata"] = {
            k: str(v) for k, v in pdf_meta.items() if v
        }

    # Try pdfplumber first (better table/layout extraction)
    try:
        with pdfplumber.open(str(file_path)) as pdf:
            tables = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                page_tables = page.extract_tables()
                if page_tables:
                    tables.extend(page_tables)
            if tables:
                metadata["tables_found"] = len(tables)
    except Exception:
        # Fallback to pypdf if pdfplumber fails (e.g. missing Ghostscript on Windows)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        metadata["parser_fallback"] = "pypdf"

    return "\n".join(text_parts), metadata


def _extract_text_from_docx(file_path: Path) -> tuple[str, dict]:
    doc = Document(str(file_path))
    metadata = {"paragraph_count": len(doc.paragraphs)}

    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                tables_text.append(" | ".join(row_text))
    if tables_text:
        text_parts.append("\n".join(tables_text))
        metadata["tables_found"] = len(doc.tables)

    return "\n".join(text_parts), metadata


def parse_resume(file_path: Path) -> ParsedResume:
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        raw_text, metadata = _extract_text_from_pdf(file_path)
        file_type = "pdf"
    elif suffix in (".docx", ".doc"):
        raw_text, metadata = _extract_text_from_docx(file_path)
        file_type = "docx"
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    sections = _segment_into_sections(raw_text)

    return ParsedResume(
        filename=file_path.name,
        file_type=file_type,
        raw_text=raw_text,
        sections=sections,
        metadata=metadata,
    )
