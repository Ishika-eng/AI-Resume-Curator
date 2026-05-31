"""Generate ATS-friendly PDF and DOCX resumes from curation results."""

import io
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    HRFlowable,
)

from app.models.export import ExportRequest


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

# Sections to skip in export (raw header data, contact info already in header)
SKIP_SECTIONS = {"Header", "Contact", "Personal Information"}


def _clean_section_content(lines: list[str], section_title: str = "") -> list[str]:
    """Clean up content lines: merge broken fragments, remove noise."""
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Skip lines that look like raw contact info in wrong section
        if re.match(r"^[\+]?\d[\d\s\-]{7,}$", stripped):
            continue  # phone number
        if re.match(r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$", stripped):
            continue  # email

        cleaned.append(stripped)

    title_lower = section_title.lower()

    # For education/skills sections, keep lines separate (don't merge)
    if any(kw in title_lower for kw in ["education", "skills", "certif", "award"]):
        return cleaned

    # For paragraph sections (profile/summary/objective), merge all lines into one
    if any(kw in title_lower for kw in ["profile", "summary", "objective", "about"]):
        non_bullet = []
        for line in cleaned:
            if line.startswith(("-", "•", "*", "·")):
                non_bullet.append(line)
            else:
                non_bullet.append(line)
        # Merge consecutive non-bullet lines into a single paragraph
        merged = []
        current_para = []
        for line in non_bullet:
            if line.startswith(("-", "•", "*", "·")):
                if current_para:
                    merged.append(" ".join(current_para))
                    current_para = []
                merged.append(line)
            else:
                current_para.append(line)
        if current_para:
            merged.append(" ".join(current_para))
        return merged

    # For other sections, merge broken lines intelligently
    merged = []
    for line in cleaned:
        if merged and not line.startswith(("-", "•", "*", "·")):
            prev = merged[-1]
            # If previous line doesn't end with punctuation and current line
            # starts with lowercase or looks like a continuation
            if prev and not prev[-1] in ".!?:,;" and len(line) < 40:
                if line[0].islower() or (not line[0].isupper() and not line[0].isdigit()):
                    merged[-1] = prev + " " + line
                    continue
        merged.append(line)

    return merged


def _get_exportable_sections(req: ExportRequest) -> list[dict]:
    """Get sections to export, skipping Header and cleaning content."""
    sections = []
    cr = req.curation_result

    for section in cr.curated_sections:
        # Skip raw header/contact sections
        if section.title in SKIP_SECTIONS:
            continue

        cleaned_content = _clean_section_content(section.content, section.title)
        if cleaned_content:
            sections.append({
                "title": section.title,
                "content": cleaned_content,
            })

    return sections


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def _set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.size = Pt(size)
    run.font.name = name
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_section_divider(doc):
    """Add a thin horizontal line after section heading."""
    border_para = doc.add_paragraph()
    border_para.paragraph_format.space_before = Pt(0)
    border_para.paragraph_format.space_after = Pt(4)
    pPr = border_para._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): "333333",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def export_docx(req: ExportRequest) -> bytes:
    doc = Document()

    # Slim margins for ATS
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    # --- Header: candidate name ---
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(req.candidate_name.upper())
    _set_run_font(name_run, size=18, bold=True, color=(33, 33, 33))

    # --- Contact line ---
    contact_parts = []
    if req.candidate_email:
        contact_parts.append(req.candidate_email)
    if req.candidate_phone:
        contact_parts.append(req.candidate_phone)
    if req.candidate_location:
        contact_parts.append(req.candidate_location)
    if req.candidate_linkedin:
        contact_parts.append(req.candidate_linkedin)
    if req.candidate_github:
        contact_parts.append(req.candidate_github)

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run("  |  ".join(contact_parts))
        _set_run_font(contact_run, size=9, color=(100, 100, 100))

    # --- Resume sections ---
    sections = _get_exportable_sections(req)

    for sect in sections:
        # Section heading
        doc.add_paragraph()  # spacer
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(sect["title"].upper())
        _set_run_font(heading_run, size=12, bold=True, color=(33, 33, 33))

        # Divider line
        _add_section_divider(doc)

        # Section content
        for line in sect["content"]:
            content_para = doc.add_paragraph()
            stripped = line.strip()

            # Bullet point style
            if stripped.startswith(("- ", "• ", "* ", "· ")):
                content_para.style = "List Bullet"
                text = stripped.lstrip("-•*· ").strip()
            else:
                text = stripped

            content_run = content_para.add_run(text)
            _set_run_font(content_run, size=11)

    # --- Missing skills section (if any) ---
    if req.curation_result.missing_skills:
        doc.add_paragraph()
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run("ADDITIONAL SKILLS (RECOMMENDED TO ADD)")
        _set_run_font(heading_run, size=12, bold=True, color=(33, 33, 33))
        _add_section_divider(doc)

        skills_text = ", ".join(ms.name for ms in req.curation_result.missing_skills)
        skills_para = doc.add_paragraph()
        skills_run = skills_para.add_run(skills_text)
        _set_run_font(skills_run, size=11)

    # Write to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

def export_pdf(req: ExportRequest) -> bytes:
    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch,
    )

    styles = getSampleStyleSheet()

    # Custom styles — ATS-friendly (no fancy graphics, plain text)
    name_style = ParagraphStyle(
        "ResumeName",
        parent=styles["Title"],
        fontSize=18,
        leading=22,
        alignment=1,  # center
        spaceAfter=2,
        textColor=HexColor("#212121"),
        fontName="Helvetica-Bold",
    )

    contact_style = ParagraphStyle(
        "ResumeContact",
        parent=styles["Normal"],
        fontSize=9,
        leading=12,
        alignment=1,
        spaceAfter=12,
        textColor=HexColor("#646464"),
        fontName="Helvetica",
    )

    heading_style = ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading2"],
        fontSize=12,
        leading=16,
        spaceBefore=14,
        spaceAfter=2,
        textColor=HexColor("#212121"),
        fontName="Helvetica-Bold",
    )

    body_style = ParagraphStyle(
        "ResumeBody",
        parent=styles["Normal"],
        fontSize=10.5,
        leading=14,
        spaceBefore=1,
        spaceAfter=2,
        textColor=HexColor("#333333"),
        fontName="Helvetica",
    )

    bullet_style = ParagraphStyle(
        "ResumeBullet",
        parent=body_style,
        leftIndent=14,
        bulletIndent=0,
        bulletFontName="Helvetica",
        bulletFontSize=10,
    )

    elements = []

    # --- Name ---
    elements.append(Paragraph(req.candidate_name.upper(), name_style))

    # --- Contact ---
    contact_parts = []
    if req.candidate_email:
        contact_parts.append(req.candidate_email)
    if req.candidate_phone:
        contact_parts.append(req.candidate_phone)
    if req.candidate_location:
        contact_parts.append(req.candidate_location)
    if req.candidate_linkedin:
        contact_parts.append(req.candidate_linkedin)
    if req.candidate_github:
        contact_parts.append(req.candidate_github)

    if contact_parts:
        elements.append(Paragraph("  |  ".join(contact_parts), contact_style))

    # --- Resume sections ---
    sections = _get_exportable_sections(req)

    for sect in sections:
        elements.append(Paragraph(sect["title"].upper(), heading_style))
        elements.append(HRFlowable(
            width="100%", thickness=0.8,
            color=HexColor("#333333"),
            spaceAfter=6, spaceBefore=0,
        ))

        for line in sect["content"]:
            stripped = line.strip()
            if stripped.startswith(("- ", "• ", "* ", "· ")):
                text = stripped.lstrip("-•*· ").strip()
                elements.append(Paragraph(f"&bull;  {text}", bullet_style))
            else:
                elements.append(Paragraph(stripped, body_style))

    # --- Missing skills ---
    if req.curation_result.missing_skills:
        elements.append(Paragraph("ADDITIONAL SKILLS (RECOMMENDED TO ADD)", heading_style))
        elements.append(HRFlowable(
            width="100%", thickness=0.8,
            color=HexColor("#333333"),
            spaceAfter=6, spaceBefore=0,
        ))
        skills_text = ", ".join(ms.name for ms in req.curation_result.missing_skills)
        elements.append(Paragraph(skills_text, body_style))

    doc.build(elements)
    buf.seek(0)
    return buf.read()
